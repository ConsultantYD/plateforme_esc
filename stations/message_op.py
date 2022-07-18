################################################################################
### 1. HEADER                                                                ###
################################################################################
# -*- coding: utf-8 -*-

from re import template
from docx import Document
import asyncio
from env_canada import ECWeather
import datetime as dt
import streamlit as st
import base64
from os.path import join

from email import generator
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

"""
@authors: Ysael Desage
"""

template_file_path = 'mo_template.docx'

def create_message(attachment):

    msg = MIMEMultipart()
    msg["To"] = "mrscqbc@dfo-mpo.gc.ca"
    msg["From"] = "undefined"
    msg["CC"] = "xca-montrealops@dfo-mpo.gc.ca"
    msg["Subject"] = "09 - GC1205 - Début des opérations 1600Z"

    body = ""

    msg.attach(MIMEText(body, "plain"))

    for path in [attachment]:
        part = MIMEBase('application', "octet-stream")
        with open(path, 'rb') as file:
            part.set_payload(file.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition',
                        'attachment', filename=attachment)
        msg.attach(part)

    return msg

def write_eml_file(msg, filename):

    with open(filename, 'w') as file:
        emlGenerator = generator.Generator(file)
        emlGenerator.flatten(msg)

def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)

def write(**kwargs):

    st.title("Messages Opérationnels")
    img_col,_,_ = st.columns((1,4,4))
    img_col.image(join("stations",'visual_mo.png'))

    # DATA ACQUISITION
    message_types = [
        "Message de début des opérations",
        "Message de fin des opérations",
        "Message de départ en patrouille",
        "Message de retour de patrouille"
    ]

    message_to = {
        "Message de début des opérations": "MRSCQBC@dfo-mpo.gc.ca",
        "Message de fin des opérations": "MRSCQBC@dfo-mpo.gc.ca",
        "Message de départ en patrouille": "Mrscq@sarnet.dnd.ca",
        "Message de retour de patrouille": "Mrscq@sarnet.dnd.ca"
    }

    message_cc = {
        "Message de début des opérations": "XCA-MontrealOps@dfo-mpo.gc.ca",
        "Message de fin des opérations": None,
        "Message de départ en patrouille": None,
        "Message de retour de patrouille": None
    }

    message_object = {
        "Message de début des opérations": "09 - GC1205 - Début des opérations 1600Z",
        "Message de fin des opérations": "09 - GC1205 - Fin des opérations 2400Z",
        "Message de départ en patrouille": "09 - GC1205 - Départ en patrouille secteur _________; __________Z",
        "Message de retour de patrouille": "09 - GC1205 - Retour de patrouille"
    }

    msg_type = st.selectbox('Choisir type de message', message_types)
    st.header(msg_type)

    col1, col2 = st.columns((1, 10))
    with col2:
        if message_to[msg_type] is not None:
            st.subheader("À")
            st.write(message_to[msg_type])
        if message_cc[msg_type] is not None:
            st.subheader("CC")
            st.write(message_cc[msg_type])
        if message_object[msg_type] is not None:
            st.subheader("Objet")
            st.write(message_object[msg_type])
        if msg_type == "Message de début des opérations":
            st.subheader("Pièces Jointes")
            mois = {
                1: "Janvier",
                2: "Février",
                3: "Mars",
                4: "Avril",
                5: "Mai",
                6: "Juin",
                7: "Juillet",
                8: "Aout",
                9: "Septembre",
                10: "Octobre",
                11: "Novembre",
                12: "Décembre"
            }

            now = dt.datetime.now()
            day = str(now.day)
            month = mois[now.month]
            year = "2022"
            date_title = day + ' ' + month + ' ' + year

            output_file_path = 'GC 1205 - Message opérationnel - ' + date_title + ' .docx'

            ec_fr = ECWeather(coordinates=(
                46.046528, -73.116527), language='french')

            asyncio.run(ec_fr.update())

            wind_sp = str(ec_fr.conditions['wind_speed']['value'])
            wind_gust = str(ec_fr.conditions['wind_gust']['value'])
            if wind_gust == "None":
                wind_gust = wind_sp
            temperature = str(ec_fr.conditions['temperature']['value'])

            variables = {
                "${DATE}": now.strftime("%Y-%m-%d"),
                "${WIND_DIR}": ec_fr.conditions['wind_dir']['value'],
                "${WIND_SP}": wind_sp,
                "${WIND_GUST}": wind_gust,
                "${VISIBILITY}": "10",
                "${TEMP}": temperature,
            }

            template_document = Document(template_file_path)

            for variable_key, variable_value in variables.items():
                for paragraph in template_document.paragraphs:
                    replace_text_in_paragraph(
                        paragraph, variable_key, variable_value)

                for table in template_document.tables:
                    for col in table.columns:
                        for cell in col.cells:
                            for paragraph in cell.paragraphs:
                                replace_text_in_paragraph(
                                    paragraph, variable_key, variable_value)

            template_document.save(output_file_path)

            # EMAIL FILE
            #filename = "message_operationnel.emltpl"
            #msg = create_message(output_file_path)
            #write_eml_file(msg, filename)

            # DOWNLOAD BUTTON

            # Remove this to use email file
            filename = output_file_path

            # Load locally for download button
            with open(filename, 'rb') as f:
                st.download_button("Telecharger", f, file_name=filename)

        st.subheader("Signature")
        st.write(
            "Signature avec le nom et le poste de la personne qui envoie le courriel.")

################################################################################
### X. END OF CODE                                                           ###
################################################################################
